"""Extracción de texto de archivos Word (.docx)."""

from __future__ import annotations

from io import BytesIO

from docx import Document


def extract_text(data: bytes) -> str:
    """Extrae párrafos y tablas de un .docx en memoria."""
    doc = Document(BytesIO(data))

    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)
